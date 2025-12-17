from docx import Document

# Read the Word document
doc = Document('assets/Fall25Notes.docx')

print(f"Reading document...")
print(f"Total paragraphs: {len(doc.paragraphs)}")

# Extract all text
all_text = []
for para in doc.paragraphs:
    text = para.text.strip()
    if text:  # Only add non-empty paragraphs
        all_text.append(text)

# Join with double newlines to preserve paragraph structure
full_text = '\n\n'.join(all_text)

# Save to text file
with open('assets/class_notes.txt', 'w', encoding='utf-8') as f:
    f.write(full_text)

print(f"✓ Conversion complete!")
print(f"Total characters: {len(full_text)}")
print(f"Saved to: assets/class_notes.txt")

# Preview first 500 characters
print(f"\nPreview of extracted text:")
print("="*60)
print(full_text[:500])
print("="*60)